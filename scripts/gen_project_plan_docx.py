"""Generate 项目计划书.docx (one-off script)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main() -> None:
    # scripts/gen_project_plan_docx.py -> repo root
    root = Path(__file__).resolve().parents[1]
    out = root / "项目计划书.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("论文润色与运营控制台项目计划书")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"

    doc.add_paragraph("文档版本：V1.0")
    doc.add_paragraph("编制日期：2026年5月")
    doc.add_paragraph("项目名称：downAiGC / Paper Polish（论文润色与 AIGC 运营一体化平台）")
    doc.add_paragraph()

    sections: list[tuple[str, str]] = [
        (
            "一、项目概述",
            """本项目建设一套面向科研与学术写作场景的「论文段落润色」Web 控制台及配套后端服务，并集成用户账户、用量与改写字数、广告激励、运营兑换码、用户反馈与管理员后台等能力。目标是在可控成本下提供稳定、可审计的润色任务流水线，支撑产品化运营与后续模型能力升级。""",
        ),
        (
            "二、建设背景与目标",
            """2.1 背景\n学术写作用户对中英文论文的段落级润色、一致性校对与历史追溯有持续需求；同时商业化需要会员/改写字数、广告曝光兑换、兑换码拉新及运营数据看板。\n\n2.2 总体目标\n（1）提供可登录、可追踪的润色任务工作台与历史记录；\n（2）后端统一鉴权、配额与任务持久化，支持多环境部署与数据库迁移；\n（3）建立管理端：用户与封禁、反馈处理、兑换码发放、运营概览；\n（4）反馈与客服链路支持图文，提升问题定位效率；\n（5）为后续接入更多模型供应商、计费精细化与小程序端扩展预留接口与数据结构。\n\n2.3 阶段目标（当前仓库已覆盖）\nMVP：润色主流程 + 认证 + 管理端基础能力 + 反馈图文 + 兑换码/余额相关数据模型与页面。""",
        ),
        (
            "三、建设范围与交付物",
            """3.1 范围内\n• 用户端：注册/登录、控制台导航、润色工作台、历史任务、设置、体验反馈（含富文本/图片）、钱包/改写字数相关页面（以仓库实现为准）。\n• 管理端：管理员登录、数据概览、用户列表与封禁、反馈列表与详情处理、兑换码生成与管理。\n• 后端：REST API、JWT 鉴权、PostgreSQL 持久化、Alembic 迁移、静态资源与反馈图片上传、广告观看票据等业务模块（以代码为准）。\n• 配套：环境变量示例、数据库迁移脚本、前端构建与本地开发说明。\n\n3.2 范围外（可列入二期）\n• 第三方支付自动对账、发票系统；\n• 多租户 SaaS 隔离与细粒度 RBAC；\n• 全自动论文全文一键排版/查重等独立产品线（可作为扩展包规划）。\n\n3.3 交付物清单\n《需求说明》（隐含于 Issue/PR 与代码注释）、《接口说明》（OpenAPI/代码路由）、《部署与迁移说明》（.env.example + Alembic）、可运行前后端工程、本《项目计划书》。""",
        ),
        (
            "四、技术架构",
            """4.1 总体架构\n浏览器（React + Vite + Ant Design）通过 HTTPS 访问 FastAPI 应用层；ORM 使用 SQLAlchemy，数据库为 PostgreSQL；静态文件由后端挂载目录提供；管理端与普通控制台共用 API 基址，通过路由与管理员鉴权区分。\n\n4.2 关键技术选型\n前端：TypeScript、React 18、React Router 6、Ant Design 5；\n后端：Python 3、FastAPI、Pydantic、SQLAlchemy 2、Alembic、python-jose、passlib；\n数据库：PostgreSQL（生产推荐）；\n部署：Uvicorn 进程 + 反向代理（Nginx/Caddy），环境变量管理密钥与 CORS。\n\n4.3 非功能约束\n安全：密码哈希、JWT 过期策略、管理员接口强鉴权、反馈图片路径白名单与 HTML 消毒；\n性能：列表分页、大文本与上传大小限制；\n可维护：分层（models / repositories / 路由）、迁移可重复执行策略。""",
        ),
        (
            "五、功能模块规划",
            """5.1 用户与认证：邮箱注册登录、验证码流程（若启用）、账号封禁状态、个人资料。\n5.2 润色任务：任务创建、段落拆分与逐段润色、模型调用与错误处理、导出与历史查询。\n5.3 用量与激励：字数/配额、改写字数或广告观看兑换（以 ad_watch 模块为准）、前端状态同步。\n5.4 钱包与兑换：用户余额/改写字数展示、兑换码核销与后台批量生成。\n5.5 反馈与客服：用户提交反馈、管理员回复（HTML/图文）、状态流转（待处理/处理中/已关闭）。\n5.6 管理后台：运营看板、用户管理、反馈处理、兑换码管理。\n5.7 扩展端：仓库内「论文润色」Uni-app 工程可作为微信小程序等渠道扩展（单独排期联调）。""",
        ),
        (
            "六、里程碑与进度计划（建议）",
            """阶段 A（2–3 周）：核心润色链路打通，任务与历史可用，基础认证与部署文档。\n阶段 B（2 周）：管理端概览与用户管理、反馈列表与详情、图片上传与展示闭环。\n阶段 C（2 周）：兑换码与余额/改写字数联动、广告观看票据联调、前端体验与国际化收尾。\n阶段 D（持续）：监控与日志、压测、安全审计、模型供应商切换与降级策略。\n\n注：以上为建议排期，可按团队人力压缩或并行。""",
        ),
        (
            "七、组织与职责",
            """• 产品经理：范围优先级、验收标准、与论文场景用户访谈；\n• 全栈/后端：API、数据库、迁移、安全与部署；\n• 前端：控制台与管理端 UI/UX、联调与性能；\n• 测试：用例、回归清单、关键路径自动化（可选）；\n• 运维：环境、备份、密钥与 CORS 配置。""",
        ),
        (
            "八、风险与应对",
            """模型供应商不稳定或涨价：抽象模型调用层，支持多密钥与降级模型。\n大模型输出质量波动：段落级重试、人工反馈闭环、提示词版本管理。\n用户上传恶意文件：类型校验、大小限制、路径白名单、HTML 消毒。\n数据丢失：定期备份、迁移前演练、重要操作审计日志（按需加表）。\n合规与版权：用户协议中明确 AI 辅助边界，禁止上传涉密未授权全文（产品策略）。""",
        ),
        (
            "九、质量目标与验收",
            """• 功能：主流程无阻塞缺陷；管理端权限正确；反馈图文端到端可用。\n• 性能：常用列表接口响应在可接受范围（具体 SLA 由运维测定）。\n• 安全：无硬编码生产密钥；HTTPS；管理员与普通用户权限隔离。\n• 验收方式：测试用例 + 演示环境走查 + 产品签字确认。""",
        ),
        (
            "十、后续演进路线",
            """• 计费：对接支付、订单与对账报表；\n• 协作：团队空间、共享任务与评论；\n• 模型：本地合并部署与云端 API 双模式统一配置中心；\n• 渠道：小程序与 Web 账号体系统一（OAuth2/手机号等按合规选型）；\n• 数据：埋点与漏斗分析、管理员导出 CSV。""",
        ),
    ]

    for title, body in sections:
        h = doc.add_paragraph()
        r = h.add_run(title)
        r.bold = True
        r.font.size = Pt(14)
        r.font.name = "黑体"
        for line in body.strip().split("\n"):
            doc.add_paragraph(line.strip())

    doc.add_paragraph()
    end = doc.add_paragraph("—— 正文结束 ——")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out))
    print(out)


if __name__ == "__main__":
    main()
